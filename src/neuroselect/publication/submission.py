"""Typed, fail-closed journal submission package assembly."""

from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
import urllib.request
import zipfile
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Literal

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pydantic import BaseModel, ConfigDict, Field, model_validator

from neuroselect.publication.archival import (
    DEFAULT_RELEASE_CONFIG,
    canonical_json,
    load_release_spec,
    publication_release_gates,
    sha256_bytes,
)
from neuroselect.publication.latex import LatexManuscript, _latex_escape, render_latex_manuscript
from neuroselect.publication.manuscript import (
    DEFAULT_MANUSCRIPT_CONFIG,
    ManuscriptSpec,
    load_manuscript_spec,
    verify_manuscript_inputs,
)

DEFAULT_SUBMISSION_CONFIG = Path("configs/publication/submission_v1.yaml")
SPRINGER_TEMPLATE_URL = (
    "https://cms-resources.apps.public.k8s.springernature.io/"
    "springer-cms/rest/v1/content/18782940/data/v12"
)
SPRINGER_TEMPLATE_SHA256 = "812e76dcaa9c28dc1bff1fb6065d51729b67d4ea140552a05088317414a3ecae"
JournalId = Literal["rbet", "neuroinformatics"]
GateStatus = Literal["pending", "satisfied", "not_applicable"]
_ORCID = re.compile(r"^https://orcid\.org/[0-9]{4}-[0-9]{4}-[0-9]{4}-[0-9]{3}[0-9X]$")
_WORD = re.compile(r"\b[\w'-]+\b")
_DIRECT_IDENTIFIER = re.compile(
    r"Nick\s+Somitsch|NickSomitsch|@uibk\.ac\.at|@gmail\.com|orcid\.org|"
    r"github\.com/[^\s/]+/|10\.5281/zenodo\.[0-9]+|\bGitHub\b|\bZenodo\b|\bv0\.1\.0\b",
    re.IGNORECASE,
)


class SubmissionGate(BaseModel):
    """Non-private record of one external decision gate."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    status: GateStatus = "pending"
    checked_on: date | None = None
    public_note: str = Field(min_length=1, max_length=300)

    @model_validator(mode="after")
    def validate_date(self) -> SubmissionGate:
        if self.status == "satisfied" and self.checked_on is None:
            raise ValueError("satisfied submission gates require a checked_on date")
        return self


class SubmissionAuthor(BaseModel):
    """Final author metadata; unresolved values remain null and block submission."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    name: str
    orcid: str | None = None
    affiliation: str | None = None
    email: str | None = None
    corresponding_author: bool = True
    contributions: str | None = None

    @model_validator(mode="after")
    def validate_orcid(self) -> SubmissionAuthor:
        if self.orcid is not None and not _ORCID.fullmatch(self.orcid):
            raise ValueError("author ORCID must use its full https://orcid.org URL")
        return self


class SubmissionDeclarations(BaseModel):
    """Author-supplied declarations; null values are never inferred."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    funding: str | None = None
    competing_interests: str | None = None
    ethics: str | None = None
    consent_to_participate: str = (
        "Not applicable; no participant was recruited and nobody used NeuroSelect."
    )
    consent_for_publication: str = (
        "Not applicable; no participant-authored communication is reported."
    )
    ai_use: str = (
        "OpenAI Codex assisted with software development, analysis workflows, and manuscript "
        "preparation. Nick Somitsch verified the resulting code, analyses, and text and retains "
        "full responsibility for the work."
    )


class SubmissionSpec(BaseModel):
    """Journal routing, author metadata, declarations, and validation rules."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    schema_version: Literal["1.0"] = "1.0"
    submission_id: Literal["neuroselect-journal-submission-v1"]
    manuscript_config: Path = DEFAULT_MANUSCRIPT_CONFIG
    manuscript_artifacts: Path
    release_config: Path = DEFAULT_RELEASE_CONFIG
    reviewer_archive_url: str | None = None
    author: SubmissionAuthor
    declarations: SubmissionDeclarations
    keywords: tuple[str, ...] = Field(min_length=4, max_length=6)
    neuroinformatics_abstract: Path
    gates: dict[str, SubmissionGate]
    common_required_gates: tuple[str, ...]
    rbet_required_gates: tuple[str, ...]

    @model_validator(mode="after")
    def validate_gates(self) -> SubmissionSpec:
        required = set(self.common_required_gates) | set(self.rbet_required_gates)
        missing = required - set(self.gates)
        if missing:
            raise ValueError(f"submission config is missing gates: {sorted(missing)}")
        if len(set(self.keywords)) != len(self.keywords):
            raise ValueError("submission keywords must be unique")
        return self


class SubmissionFile(BaseModel):
    """One journal portal deliverable."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    filename: str
    portal_designation: str
    size: int = Field(ge=1)
    sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    license: str
    public_status: Literal["public", "private-review", "editor-only"]


class SubmissionInventory(BaseModel):
    """Machine-readable journal package inventory."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    schema_version: Literal["1.0"] = "1.0"
    submission_id: str
    journal: JournalId
    article_type: str
    route: str
    git_revision: str = Field(pattern=r"^[0-9a-f]{40}$")
    release_tag: Literal["v0.1.0"]
    zenodo_doi: str | None = None
    submission_ready: bool
    development_preview: bool
    pending_gates: tuple[str, ...]
    abstract_word_count: int | None = None
    keyword_count: int
    files: tuple[SubmissionFile, ...] = Field(min_length=1)
    archive_filename: str
    archive_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")


def load_submission_spec(path: str | Path = DEFAULT_SUBMISSION_CONFIG) -> SubmissionSpec:
    """Load and validate the journal submission configuration."""

    payload = yaml.safe_load(Path(path).read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError("submission config must contain a YAML mapping")
    return SubmissionSpec.model_validate(payload)


def selected_journal(spec: SubmissionSpec) -> JournalId:
    """Apply the €0 route: RBET only after both institutional gates are satisfied."""

    if all(spec.gates[gate].status == "satisfied" for gate in spec.rbet_required_gates):
        return "rbet"
    return "neuroinformatics"


def _pending_gates(
    spec: SubmissionSpec, journal: JournalId, repository: Path | None = None
) -> list[str]:
    ids = list(spec.common_required_gates)
    if journal == "rbet":
        ids.extend(spec.rbet_required_gates)
    pending = [gate for gate in ids if spec.gates[gate].status != "satisfied"]
    if spec.author.orcid is None:
        pending.append("orcid_metadata")
    if spec.author.affiliation is None:
        pending.append("author_affiliation")
    if spec.author.email is None:
        pending.append("corresponding_author_email")
    if spec.author.contributions is None:
        pending.append("credit_contributions")
    if spec.declarations.funding is None:
        pending.append("funding_declaration")
    if spec.declarations.competing_interests is None:
        pending.append("competing_interests_declaration")
    if spec.declarations.ethics is None:
        pending.append("secondary_use_ethics_wording")
    release = load_release_spec(spec.release_config)
    if release.zenodo_doi is None:
        pending.append("reserved_zenodo_doi")
    if repository is not None:
        pending.extend(
            f"publication_release: {gate}"
            for gate in publication_release_gates(release, repository.resolve())
        )
    if journal == "rbet" and spec.reviewer_archive_url is None:
        pending.append("anonymous_reviewer_archive")
    return list(dict.fromkeys(pending))


def _replace_section(source: str, start: str, end: str, replacement: str) -> str:
    prefix, separator, remainder = source.partition(start)
    if not separator:
        raise ValueError(f"manuscript section is missing: {start.strip()}")
    _, separator, suffix = remainder.partition(end)
    if not separator:
        raise ValueError(f"manuscript section terminator is missing: {end.strip()}")
    return prefix + replacement.rstrip() + "\n\n" + end + suffix


def _declaration_source(spec: SubmissionSpec, release_doi: str | None) -> str:
    author_contributions = spec.author.contributions or "PENDING: confirm final CRediT roles."
    funding = spec.declarations.funding or "PENDING: author confirmation required."
    conflicts = spec.declarations.competing_interests or "PENDING: author confirmation required."
    ethics = spec.declarations.ethics or "PENDING: approved secondary-use wording required."
    doi = release_doi or "PENDING: reserve the version-specific Zenodo DOI"
    return f"""# Information Sharing Statement

NeuroSelect source code, protocols, tests, and manuscript sources are archived at the public
GitHub tag `v0.1.0` and Zenodo DOI {doi}. Study P is public, deidentified secondary data available
from PhysioNet under its source terms; raw EEG is not redistributed. Manifest-verified,
non-restricted results, tables, figures, and reports cited by this article accompany the release.
Qwen weights, caches, LoRA adapters, executable checkpoints, and private correspondence are
excluded.

# Acknowledgments

No institutional supervision, funding, ethics approval, or endorsement is claimed unless recorded
in the final author metadata. Any independent scientific review that does not meet authorship
criteria will be acknowledged here only with the reviewer's permission.

# Declarations

## Funding

{funding}

## Competing interests

{conflicts}

## Ethics approval

{ethics} No participant was recruited and nobody used NeuroSelect. Study P was analyzed as public,
deidentified secondary data.

## Consent to participate

{spec.declarations.consent_to_participate}

## Consent for publication

{spec.declarations.consent_for_publication}

## Data availability

Study P remains available from PhysioNet under its source terms. Synthetic and counterfactual
messages were not participant-authored communication and do not express participant intent.

## Code availability

The exact submitted software is identified by GitHub tag `v0.1.0` and Zenodo DOI {doi}.

## Author contributions

{author_contributions}

## Use of AI-assisted tools

{spec.declarations.ai_use}
"""


def _submission_source(spec: SubmissionSpec, journal: JournalId, manuscript: str) -> str:
    release = load_release_spec(spec.release_config)
    source = manuscript
    if journal == "neuroinformatics":
        abstract = spec.neuroinformatics_abstract.read_text(encoding="utf-8").strip()
        count = len(_WORD.findall(abstract))
        if not 150 <= count <= 250:
            raise ValueError(f"Neuroinformatics abstract must contain 150-250 words, found {count}")
        source = _replace_section(
            source, "# Abstract\n", "# 1. Introduction", "# Abstract\n\n" + abstract
        )
    results_marker = "# 3. Results"
    if results_marker not in source:
        raise ValueError("journal manuscript is missing the Results section marker")
    source = source.replace(
        results_marker,
        "## 2.9 Use of AI-assisted tools\n\n" + spec.declarations.ai_use + "\n\n" + results_marker,
        1,
    )
    source = _replace_section(
        source,
        "# Declarations\n",
        "# References",
        _declaration_source(spec, release.zenodo_doi),
    )
    if any(line.startswith("####") for line in source.splitlines()):
        raise ValueError("journal manuscript may use no more than three heading levels")
    return source


def _compile_latex(
    source: str, bibliography: str, figures: dict[str, bytes], extras: dict[str, bytes]
) -> bytes:
    tectonic = shutil.which("tectonic")
    if tectonic is None:
        raise RuntimeError("Tectonic is required to compile a journal submission")
    with tempfile.TemporaryDirectory(prefix="neuroselect-submission-") as temporary:
        workspace = Path(temporary)
        (workspace / "manuscript.tex").write_text(source, encoding="utf-8")
        (workspace / "references.bib").write_text(bibliography, encoding="utf-8")
        for name, content in {**figures, **extras}.items():
            target = workspace / name
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(content)
        result = subprocess.run(
            [tectonic, "--keep-logs", "--outdir", str(workspace), "manuscript.tex"],
            cwd=workspace,
            capture_output=True,
            text=True,
        )
        if result.returncode:
            diagnostic = "\n".join((result.stdout + result.stderr).splitlines()[-50:])
            raise RuntimeError(f"journal LaTeX compilation failed:\n{diagnostic}")
        pdf = (workspace / "manuscript.pdf").read_bytes()
    if not pdf.startswith(b"%PDF"):
        raise RuntimeError("journal compilation did not produce a PDF")
    return pdf


def _springer_template(cache: Path) -> dict[str, bytes]:
    cache.parent.mkdir(parents=True, exist_ok=True)
    if not cache.exists():
        with urllib.request.urlopen(SPRINGER_TEMPLATE_URL, timeout=60) as response:
            cache.write_bytes(response.read())
    content = cache.read_bytes()
    if sha256_bytes(content) != SPRINGER_TEMPLATE_SHA256:
        raise ValueError("official Springer Nature LaTeX template checksum changed")
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return {
            "sn-jnl.cls": archive.read("sn-article-template/sn-jnl.cls"),
            "sn-basic.bst": archive.read("sn-article-template/bst/sn-basic.bst"),
        }


def _springer_latex(
    base: LatexManuscript,
    spec: SubmissionSpec,
    manuscript_spec: ManuscriptSpec,
    abstract: str,
) -> str:
    body = base.source.split("\\begin{document}", 1)[1].rsplit("\\end{document}", 1)[0]
    body = body.split("\\maketitle", 1)[1]
    body = re.sub(r"\\begin\{abstract\}.*?\\end\{abstract\}", "", body, flags=re.DOTALL)
    body = body.replace(r"\bibliographystyle{unsrt}", "")
    body = body.replace(r"\cite{", r"\citep{")
    # Use table* to avoid sn-jnl's threeparttable wrapper while retaining the
    # journal-neutral resizebox that keeps wide tables within the page boundary.
    body = body.replace(r"\begin{table}[H]", r"\begin{table*}[p]")
    body = body.replace(r"\end{table}", r"\end{table*}")
    body = body.replace(
        r"\section{Discussion}",
        "\\clearpage\n\\section{Discussion}",
        1,
    )
    body = body.replace(
        r"\section*{Declarations}",
        r"\section{Statements and Declarations}",
        1,
    )
    abstract_match = re.search(
        r"\\begin\{abstract\}(.*?)\\end\{abstract\}", base.source, flags=re.DOTALL
    )
    if abstract_match is None:
        raise ValueError("rendered manuscript is missing its abstract")
    abstract_latex = abstract_match.group(1).strip()
    affiliation = _latex_escape(spec.author.affiliation or "PENDING affiliation")
    email = _latex_escape(spec.author.email or "pending@example.invalid")
    orcid = spec.author.orcid or "PENDING ORCID"
    author_parts = spec.author.name.rsplit(maxsplit=1)
    if len(author_parts) == 2:
        given_name, family_name = author_parts
    else:
        given_name, family_name = "", author_parts[0]
    author_name = (
        rf"\fnm{{{_latex_escape(given_name)}}} " if given_name else ""
    ) + rf"\sur{{{_latex_escape(family_name)}}}"
    return rf"""\documentclass[pdflatex,sn-basic]{{sn-jnl}}
\usepackage[T1]{{fontenc}}
\usepackage[utf8]{{inputenc}}
\usepackage{{graphicx}}
\usepackage{{amsmath,amssymb,amsfonts}}
\usepackage{{amsthm}}
\usepackage{{booktabs}}
\usepackage{{array}}
\usepackage{{float}}
\usepackage{{pdflscape}}
\usepackage{{xurl}}
\graphicspath{{{{figures/}}}}
\begin{{document}}
\title[{_latex_escape(manuscript_spec.title)}]{{{_latex_escape(manuscript_spec.title)}}}
\author*[1]{{{author_name}}}
\email{{{email}}}
\affil*[1]{{\orgname{{{affiliation}}}}}
\abstract{{{abstract_latex}}}
\keywords{{{_latex_escape(", ".join(spec.keywords))}}}
% ORCID: {orcid}
\maketitle
{body.strip()}
\end{{document}}
"""


def _anonymous_text(text: str) -> str:
    text = re.sub(r"Nick\s+Somitsch|NickSomitsch", "Anonymous", text, flags=re.IGNORECASE)
    text = re.sub(r"https://github\.com/[^\s}<\"]+", "anonymous-review-materials", text)
    text = re.sub(r"10\.5281/zenodo\.[0-9]+", "anonymous-review-record", text)
    text = re.sub(r"https://orcid\.org/[^\s}<\"]+", "anonymous-orcid", text)
    text = re.sub(r"\bGitHub\b", "anonymous repository", text)
    text = re.sub(r"\bZenodo\b", "anonymous archive", text)
    text = text.replace("v0.1.0", "review snapshot")
    return text


def _rewrite_docx(content: bytes, *, anonymize: bool) -> bytes:
    source = io.BytesIO(content)
    target = io.BytesIO()
    with zipfile.ZipFile(source) as archive, zipfile.ZipFile(target, "w") as output:
        for info in archive.infolist():
            data = archive.read(info.filename)
            if anonymize and info.filename.endswith((".xml", ".rels")):
                data = _anonymous_text(data.decode("utf-8")).encode()
            fixed = zipfile.ZipInfo(info.filename, (1980, 1, 1, 0, 0, 0))
            fixed.compress_type = zipfile.ZIP_DEFLATED
            fixed.external_attr = 0o644 << 16
            output.writestr(fixed, data)
    return target.getvalue()


def _anonymous_docx(content: bytes) -> bytes:
    return _rewrite_docx(content, anonymize=True)


def _deterministic_zip(files: dict[str, bytes]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        for name, content in sorted(files.items()):
            if Path(name).is_absolute() or ".." in Path(name).parts:
                raise ValueError(f"unsafe submission archive path: {name}")
            info = zipfile.ZipInfo(name, (1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            archive.writestr(info, content)
    return output.getvalue()


def _git_revision(repository: Path) -> str:
    return subprocess.run(
        ["git", "rev-parse", "HEAD"], cwd=repository, check=True, capture_output=True, text=True
    ).stdout.strip()


def _simple_docx(title: str, paragraphs: tuple[str, ...], *, author: str) -> bytes:
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = author
    document.core_properties.subject = "NeuroSelect journal submission"
    fixed_time = datetime(2026, 7, 31, tzinfo=UTC)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(9)
    buffer = io.BytesIO()
    document.save(buffer)
    return _rewrite_docx(buffer.getvalue(), anonymize=False)


def _file_record(
    name: str,
    content: bytes,
    designation: str,
    status: Literal["public", "private-review", "editor-only"],
) -> SubmissionFile:
    license_id = (
        "LPPL-1.0-or-later" if name.endswith(("sn-jnl.cls", "sn-basic.bst")) else "CC-BY-4.0"
    )
    return SubmissionFile(
        filename=name,
        portal_designation=designation,
        size=len(content),
        sha256=sha256_bytes(content),
        license=license_id,
        public_status=status,
    )


def _shared_supplements(manuscript_spec: ManuscriptSpec) -> dict[str, bytes]:
    files: dict[str, bytes] = {}
    display = manuscript_spec.display.path
    for path in sorted((display / "figures").glob("*.pdf")):
        files[f"figures/{path.name}"] = path.read_bytes()
    for path in sorted((display / "tables").glob("*.csv")):
        files[f"tables/{path.name}"] = path.read_bytes()
    for name in ("captions.md", "inventory.json"):
        files[f"supplement/{name}"] = (display / name).read_bytes()
    files["supplement/claim-audit.json"] = Path(
        "artifacts/publication/manuscript-v1/claim-audit.json"
    ).read_bytes()
    return files


def build_journal_submission(
    spec: SubmissionSpec,
    *,
    journal: JournalId,
    repository: Path,
    output: Path,
    allow_pending: bool = False,
    overwrite: bool = False,
) -> SubmissionInventory:
    """Build a validated RBET or Neuroinformatics submission package."""

    recommended = selected_journal(spec)
    if journal == "rbet" and recommended != "rbet" and not allow_pending:
        raise ValueError(
            "RBET is blocked until UIBK affiliation and full APC coverage are confirmed"
        )
    repository = repository.resolve()
    release = load_release_spec(spec.release_config)
    pending = _pending_gates(spec, journal, repository)
    if pending and not allow_pending:
        raise ValueError("journal submission is blocked: " + "; ".join(pending))
    if output.exists():
        if not overwrite:
            raise FileExistsError(f"submission output already exists: {output}")
        shutil.rmtree(output)
    output.mkdir(parents=True)
    manuscript_spec = load_manuscript_spec(spec.manuscript_config)
    manuscript, references, _, display_inventory, _ = verify_manuscript_inputs(manuscript_spec)
    source = _submission_source(spec, journal, manuscript)
    base = render_latex_manuscript(manuscript_spec, source, references, display_inventory)
    payloads = _shared_supplements(manuscript_spec)
    abstract_count: int | None = None
    article_type: str
    route: str
    if journal == "neuroinformatics":
        article_type = "Original Article"
        route = "subscription-no-apc"
        abstract = spec.neuroinformatics_abstract.read_text(encoding="utf-8").strip()
        abstract_count = len(_WORD.findall(abstract))
        template = _springer_template(
            Path("artifacts/cache/springer-nature-latex-template-v12.zip")
        )
        latex = _springer_latex(base, spec, manuscript_spec, abstract)
        pdf = _compile_latex(latex, base.bibliography, base.figures, template)
        payloads.update(
            {
                "manuscript/manuscript.tex": latex.encode(),
                "manuscript/manuscript.pdf": pdf,
                "manuscript/references.bib": base.bibliography.encode(),
                "manuscript/sn-jnl.cls": template["sn-jnl.cls"],
                "manuscript/sn-basic.bst": template["sn-basic.bst"],
                "manuscript/SPRINGER-TEMPLATE-SOURCE.txt": (
                    f"{SPRINGER_TEMPLATE_URL}\nSHA-256: {SPRINGER_TEMPLATE_SHA256}\n"
                ).encode(),
            }
        )
        payloads["editor/cover-letter.docx"] = _simple_docx(
            "Cover letter",
            (
                "Dear Editors,",
                "Please consider this manuscript as an Original Article in Neuroinformatics "
                "using the subscription publication route.",
                "No participant was recruited and nobody used NeuroSelect. Study P was analyzed "
                "as public, deidentified secondary data; synthetic and counterfactual messages "
                "were not participant-authored communication.",
                "All weak, null, and unfavorable findings were retained. The work does not "
                "demonstrate clinical efficacy, live communication benefit, or thought decoding. "
                "This manuscript is not under consideration elsewhere.",
            ),
            author=spec.author.name,
        )
    else:
        article_type = "Original Research"
        route = "open-access-uibk-covered"
        full_pdf = _compile_latex(base.source, base.bibliography, base.figures, {})
        anonymous_latex = _anonymous_text(base.source)
        anonymous_pdf = _compile_latex(anonymous_latex, base.bibliography, base.figures, {})
        source_docx = spec.manuscript_artifacts / manuscript_spec.output_filename
        anonymous_docx = _anonymous_docx(source_docx.read_bytes())
        title_page = _simple_docx(
            manuscript_spec.title,
            (
                spec.author.name,
                spec.author.affiliation or "PENDING affiliation",
                spec.author.orcid or "PENDING ORCID",
                spec.author.email or "PENDING corresponding-author email",
            ),
            author=spec.author.name,
        )
        cover = _simple_docx(
            "Cover letter",
            (
                "Dear Editors,",
                "Please consider this Original Research article for the Brain-Computer "
                "Interfaces and Neural Engineering section.",
                "No participant was recruited and nobody used NeuroSelect. The public GitHub "
                "repository may make author identity discoverable; the reviewer-facing files "
                "nevertheless remove direct identifiers.",
                "All weak, null, and unfavorable findings were retained. This manuscript is not "
                "under consideration elsewhere.",
            ),
            author=spec.author.name,
        )
        reviewer_files = {
            "anonymous-manuscript.docx": anonymous_docx,
            "anonymous-manuscript.pdf": anonymous_pdf,
            **{
                name: data
                for name, data in payloads.items()
                if name.startswith(("figures/", "tables/"))
            },
        }
        payloads.update(
            {
                "manuscript/complete-manuscript.tex": base.source.encode(),
                "manuscript/complete-manuscript.pdf": full_pdf,
                "manuscript/references.bib": base.bibliography.encode(),
                "review/anonymous-manuscript.tex": anonymous_latex.encode(),
                "review/anonymous-manuscript.pdf": anonymous_pdf,
                "review/anonymous-manuscript.docx": anonymous_docx,
                "review/anonymous-reviewer-materials.zip": _deterministic_zip(reviewer_files),
                "editor/title-page.docx": title_page,
                "editor/cover-letter.docx": cover,
            }
        )
    payloads["editor/author-metadata.json"] = canonical_json(spec.author.model_dump(mode="json"))
    payloads["editor/declarations.json"] = canonical_json(spec.declarations.model_dump(mode="json"))
    status_map: dict[str, tuple[str, Literal["public", "private-review", "editor-only"]]] = {
        "review/": ("Anonymous manuscript/reviewer material", "private-review"),
        "editor/": ("Editor-only metadata, title page, or cover letter", "editor-only"),
        "manuscript/": ("Main manuscript", "public"),
        "figures/": ("Figure", "public"),
        "tables/": ("Table", "public"),
        "supplement/": ("Supplementary material", "public"),
    }
    records = []
    for name, content in sorted(payloads.items()):
        designation, status = next(
            value for prefix, value in status_map.items() if name.startswith(prefix)
        )
        records.append(_file_record(name, content, designation, status))
        target = output / name
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(content)
    archive_name = f"neuroselect-{journal}-submission-v1.zip"
    provisional = SubmissionInventory(
        submission_id=spec.submission_id,
        journal=journal,
        article_type=article_type,
        route=route,
        git_revision=_git_revision(repository),
        release_tag=release.tag,
        zenodo_doi=release.zenodo_doi,
        submission_ready=not pending,
        development_preview=bool(pending),
        pending_gates=tuple(pending),
        abstract_word_count=abstract_count,
        keyword_count=len(spec.keywords),
        files=tuple(records),
        archive_filename=archive_name,
        archive_sha256="0" * 64,
    )
    inventory_bytes = canonical_json(provisional.model_dump(mode="json"))
    archive = _deterministic_zip({**payloads, "submission-inventory.json": inventory_bytes})
    inventory = provisional.model_copy(update={"archive_sha256": sha256_bytes(archive)})
    inventory_bytes = canonical_json(inventory.model_dump(mode="json"))
    archive = _deterministic_zip({**payloads, "submission-inventory.json": inventory_bytes})
    inventory = inventory.model_copy(update={"archive_sha256": sha256_bytes(archive)})
    # The inventory outside the archive is authoritative; archive checksum is not self-embedded.
    (output / "submission-inventory.json").write_bytes(
        canonical_json(inventory.model_dump(mode="json"))
    )
    (output / archive_name).write_bytes(archive)
    verify_journal_submission(
        output,
        require_ready=not allow_pending,
        repository=repository,
    )
    return inventory


def _anonymous_payload_text(path: Path) -> str:
    if path.suffix == ".docx":
        with zipfile.ZipFile(path) as archive:
            return "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.endswith((".xml", ".rels"))
            )
    if path.suffix == ".pdf" and shutil.which("pdftotext"):
        result = subprocess.run(["pdftotext", str(path), "-"], capture_output=True, text=True)
        metadata = subprocess.run(["pdfinfo", str(path)], capture_output=True, text=True).stdout
        return result.stdout + metadata
    return path.read_text(encoding="utf-8", errors="ignore")


def _verify_pdf_fonts(output: Path, inventory: SubmissionInventory) -> None:
    pdffonts = shutil.which("pdffonts")
    if pdffonts is None:
        return
    for item in inventory.files:
        if not item.filename.lower().endswith(".pdf"):
            continue
        result = subprocess.run(
            [pdffonts, str(output / item.filename)],
            capture_output=True,
            text=True,
        )
        if result.returncode:
            raise ValueError(f"invalid PDF in journal package: {item.filename}")
        for line in result.stdout.splitlines()[2:]:
            fields = line.split()
            if fields and (len(fields) < 5 or fields[-5] != "yes"):
                raise ValueError(f"unembedded font in journal PDF: {item.filename}")


def verify_journal_submission(
    output: Path, *, require_ready: bool = True, repository: Path | None = None
) -> SubmissionInventory:
    """Validate checksums, journal rules, and reviewer anonymity."""

    inventory = SubmissionInventory.model_validate_json(
        (output / "submission-inventory.json").read_text(encoding="utf-8")
    )
    if require_ready and not inventory.submission_ready:
        raise ValueError("submission package is not ready: " + "; ".join(inventory.pending_gates))
    if require_ready and inventory.zenodo_doi is None:
        raise ValueError("submission package does not identify a reserved Zenodo DOI")
    if require_ready and repository is not None:
        repository = repository.resolve()
        revision = _git_revision(repository)
        if revision != inventory.git_revision:
            raise ValueError("submission package Git revision does not match the checkout")
        tags = subprocess.run(
            ["git", "tag", "--points-at", revision],
            cwd=repository,
            check=True,
            capture_output=True,
            text=True,
        ).stdout.splitlines()
        if inventory.release_tag not in tags:
            raise ValueError("submission release tag does not point to the package commit")
        status = subprocess.run(
            ["git", "status", "--porcelain"],
            cwd=repository,
            check=True,
            capture_output=True,
            text=True,
        ).stdout
        if status:
            raise ValueError("submission verification requires a clean worktree")
    for item in inventory.files:
        content = (output / item.filename).read_bytes()
        if len(content) != item.size or sha256_bytes(content) != item.sha256:
            raise ValueError(f"submission file checksum mismatch: {item.filename}")
    archive = (output / inventory.archive_filename).read_bytes()
    if sha256_bytes(archive) != inventory.archive_sha256:
        raise ValueError("submission archive checksum mismatch")
    _verify_pdf_fonts(output, inventory)
    manuscript_source = (
        output
        / (
            "manuscript/manuscript.tex"
            if inventory.journal == "neuroinformatics"
            else "manuscript/complete-manuscript.tex"
        )
    ).read_text(encoding="utf-8")
    results_heading = re.search(r"\\section\{(?:3\.\s*)?Results\}", manuscript_source)
    if (
        "Use of AI-assisted tools" not in manuscript_source
        or results_heading is None
        or manuscript_source.index("Use of AI-assisted tools") > results_heading.start()
    ):
        raise ValueError("AI-use disclosure must appear in the manuscript Methods section")
    if inventory.journal == "neuroinformatics":
        if inventory.article_type != "Original Article" or inventory.route != "subscription-no-apc":
            raise ValueError("Neuroinformatics must use the €0 Original Article subscription route")
        if inventory.abstract_word_count is None or not 150 <= inventory.abstract_word_count <= 250:
            raise ValueError("Neuroinformatics abstract word count is invalid")
        if not 4 <= inventory.keyword_count <= 6:
            raise ValueError("Neuroinformatics keyword count is invalid")
        source = manuscript_source
        required = (
            r"\documentclass[pdflatex,sn-basic]{sn-jnl}",
            r"\bibliography{references}",
            "Information Sharing Statement",
            "Acknowledgments",
            "Statements and Declarations",
            "Use of AI-assisted tools",
        )
        if any(value not in source for value in required):
            raise ValueError("Neuroinformatics LaTeX source is missing a required element")
        if source.index("Information Sharing Statement") > source.index("Acknowledgments"):
            raise ValueError("Information Sharing Statement must precede acknowledgments")
    else:
        if inventory.article_type != "Original Research":
            raise ValueError("RBET package must use Original Research")
        for item in inventory.files:
            if item.public_status == "private-review":
                text = _anonymous_payload_text(output / item.filename)
                if _DIRECT_IDENTIFIER.search(text):
                    raise ValueError(f"direct identifier in reviewer-facing file: {item.filename}")
    return inventory
