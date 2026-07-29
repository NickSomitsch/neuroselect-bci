"""Assemble the verified NeuroSelect manuscript as Markdown, DOCX, LaTeX, and PDF."""

from __future__ import annotations

import argparse
import csv
import hashlib
import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import TYPE_CHECKING, Any, cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from neuroselect.publication import read_publication_display
from neuroselect.publication.latex import LatexManuscript, render_latex_manuscript
from neuroselect.publication.manuscript import (
    DEFAULT_MANUSCRIPT_CONFIG,
    ManuscriptReference,
    ManuscriptSpec,
    load_manuscript_spec,
    replace_citations,
    verify_manuscript_inputs,
    write_manuscript_bundle,
)

if TYPE_CHECKING:
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

BLUE = "0072B2"
DARK_BLUE = "1F4E79"
BLACK = "222222"
GRAY = "5F6B76"
LIGHT_GRAY = "F4F6F9"
PALE_BLUE = "EAF3F8"
_MARKER_PATTERN = re.compile(r"^\{\{(figure|table):([a-z0-9-]+)\}\}$")
_BOLD_PATTERN = re.compile(r"(\*\*.+?\*\*)")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", type=Path, default=DEFAULT_MANUSCRIPT_CONFIG)
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("artifacts/publication/manuscript-v1"),
    )
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument(
        "--allow-dirty",
        action="store_true",
        help="Permit development assembly; the inventory will remain non-ready.",
    )
    parser.add_argument(
        "--sync-latex-source",
        action="store_true",
        help="Rewrite the tracked LaTeX source and bibliography from verified inputs.",
    )
    return parser.parse_args()


def git_state() -> tuple[str, str | None]:
    revision = subprocess.run(
        ["git", "rev-parse", "HEAD"], check=True, capture_output=True, text=True
    ).stdout.strip()
    status = subprocess.run(
        ["git", "status", "--porcelain", "--untracked-files=normal"],
        check=True,
        capture_output=True,
        text=True,
    ).stdout
    if not status:
        return revision, None
    digest = hashlib.sha256(
        subprocess.run(["git", "diff", "--binary", "HEAD"], check=True, capture_output=True).stdout
    )
    untracked = subprocess.run(
        ["git", "ls-files", "--others", "--exclude-standard", "-z"],
        check=True,
        capture_output=True,
    ).stdout.split(b"\0")
    for raw_path in sorted(path for path in untracked if path):
        path = Path(raw_path.decode())
        digest.update(raw_path)
        digest.update(b"\0")
        digest.update(path.read_bytes())
        digest.update(b"\0")
    return revision, digest.hexdigest()


def _set_cell_shading(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: _Cell, *, top: int, start: int, bottom: int, end: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row: object) -> None:
    tr_pr = cast("Any", row)._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _prevent_row_split(row: object) -> None:
    tr_pr = cast("Any", row)._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_table_width(table: Table, width_twips: int = 9360) -> None:
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_cell_width(cell: _Cell, width_twips: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def _column_widths(headers: list[str], rows: list[list[str]], total: int = 9360) -> list[int]:
    weights: list[float] = []
    for index, header in enumerate(headers):
        values = [header, *(row[index] for row in rows)]
        longest = min(max(len(value) for value in values), 42)
        weights.append(max(8.0, longest**0.72))
    scale = total / sum(weights)
    widths = [max(520, round(weight * scale)) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def _set_run_font(run: Run, size: float, *, bold: bool = False, color: str = BLACK) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), "Calibri")


def _add_page_break(document: DocumentObject) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_field(run: Run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field = OxmlElement("w:instrText")
    field.set(qn("xml:space"), "preserve")
    field.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, field, separate, text, end):
        run._r.append(element)


def _configure_document(document: DocumentObject, spec: ManuscriptSpec) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_settings = (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    )
    for name, size, color, before, after in heading_settings:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("Manuscript Caption", "Manuscript Note", "Reference Entry"):
        if style_name not in styles:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Manuscript Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.bold = False
    caption.font.color.rgb = RGBColor.from_string(BLACK)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True
    note = styles["Manuscript Note"]
    note.font.name = "Calibri"
    note.font.size = Pt(8)
    note.font.color.rgb = RGBColor.from_string(GRAY)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(8)
    reference = styles["Reference Entry"]
    reference.font.name = "Calibri"
    reference.font.size = Pt(9)
    reference.paragraph_format.left_indent = Inches(0.28)
    reference.paragraph_format.first_line_indent = Inches(-0.28)
    reference.paragraph_format.space_after = Pt(5)
    reference.paragraph_format.line_spacing = 1.15

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = "NEUROSELECT  ·  OFFLINE COMPUTATIONAL ORIGINAL RESEARCH"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if paragraph.runs:
        _set_run_font(paragraph.runs[0], 8, bold=True, color=GRAY)
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_run_font(run, 8, color=GRAY)
    _add_field(run, "PAGE")

    properties = document.core_properties
    properties.title = spec.title
    properties.author = spec.author
    properties.subject = "Journal-neutral offline computational Original Research manuscript"
    properties.keywords = "BCI, P300, language prediction, personalization, reproducibility"
    properties.comments = (
        "Assembled from checksum-verified NeuroSelect publication evidence. "
        "External submission gates remain separate."
    )


def _add_rule(paragraph: Paragraph, *, color: str = BLUE, size: str = "14") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _word_count(source: str) -> int:
    stripped = re.sub(r"\{\{[^}]+\}\}", "", source)
    stripped = re.sub(r"\[@[^\]]+\]", "", stripped)
    stripped = re.sub(r"[#*_`\\()]", "", stripped)
    return len(re.findall(r"\b[\w'-]+\b", stripped))


def _add_cover(document: DocumentObject, spec: ManuscriptSpec, manuscript_source: str) -> None:
    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_before = Pt(26)
    eyebrow.paragraph_format.space_after = Pt(20)
    run = eyebrow.add_run("ORIGINAL RESEARCH  /  OFFLINE COMPUTATIONAL STUDY")
    _set_run_font(run, 9, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(spec.title)
    _set_run_font(run, 25, bold=True, color=BLACK)
    _add_rule(title)

    author = document.add_paragraph()
    author.paragraph_format.space_before = Pt(8)
    author.paragraph_format.space_after = Pt(3)
    run = author.add_run(spec.author)
    _set_run_font(run, 12, bold=True, color=DARK_BLUE)

    metadata = document.add_paragraph()
    metadata.paragraph_format.space_after = Pt(20)
    run = metadata.add_run(
        f"Journal-neutral manuscript  ·  {spec.assembled_at.date().isoformat()}  ·  "
        f"{_word_count(manuscript_source):,} source words"
    )
    _set_run_font(run, 9, color=GRAY)

    scope = document.add_table(rows=3, cols=2)
    scope.alignment = WD_TABLE_ALIGNMENT.LEFT
    scope.autofit = False
    _set_table_width(scope, 7900)
    scope_data = (
        ("Evidence", "Synthetic language · original-task public EEG · counterfactual replay"),
        (
            "Primary boundary",
            "No participant used NeuroSelect; remapped phrases are not user intent",
        ),
        ("Assembly state", "Verified manuscript draft; external submission gates remain separate"),
    )
    for row_index, (label, value) in enumerate(scope_data):
        cells = scope.rows[row_index].cells
        _set_table_cell_width(cells[0], 1600)
        _set_table_cell_width(cells[1], 6300)
        _set_cell_shading(cells[0], PALE_BLUE)
        for cell in cells:
            _set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].text = label
        cells[1].text = value
        _set_run_font(cells[0].paragraphs[0].runs[0], 8.5, bold=True, color=DARK_BLUE)
        _set_run_font(cells[1].paragraphs[0].runs[0], 8.5, color=BLACK)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Submission note. Affiliation, institutional email, ORCID, funding, competing interests, "
        "and final CRediT metadata must be confirmed before journal submission."
    )
    _set_run_font(run, 8.5, color=GRAY)
    _add_page_break(document)


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    clean = text.replace(r"\(", "").replace(r"\)", "")
    for chunk in _BOLD_PATTERN.split(clean):
        if not chunk:
            continue
        bold = chunk.startswith("**") and chunk.endswith("**")
        content = chunk[2:-2] if bold else chunk
        run = paragraph.add_run(content)
        _set_run_font(run, 11, bold=bold)


def _add_body_paragraph(document: DocumentObject, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    _add_inline_runs(paragraph, text)


def _table_label(item_id: str) -> str:
    match = re.match(r"table-([0-9]+[a-z]?)", item_id)
    if match is None:
        raise ValueError(f"invalid table ID: {item_id}")
    return f"Table {match.group(1)}"


def _figure_label(item_id: str) -> str:
    match = re.match(r"figure-([0-9]+)", item_id)
    if match is None:
        raise ValueError(f"invalid figure ID: {item_id}")
    return f"Figure {match.group(1)}"


def _add_table(document: DocumentObject, display_path: Path, item: object) -> None:
    item_id = cast("Any", item).item_id
    csv_path = display_path / next(
        file for file in cast("Any", item).files if file.endswith(".csv")
    )
    with csv_path.open(encoding="utf-8", newline="") as source_file:
        reader = csv.reader(source_file)
        rows = list(reader)
    headers = rows[0]
    body = rows[1:]
    label = _table_label(item_id)
    caption = document.add_paragraph(style="Manuscript Caption")
    first = caption.add_run(f"{label}. {cast('Any', item).title}. ")
    _set_run_font(first, 9, bold=True)
    second = caption.add_run(cast("Any", item).caption)
    _set_run_font(second, 9)

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_width(table)
    widths = _column_widths(headers, body)
    font_size = 6.4 if len(headers) >= 9 else 7.0 if len(headers) >= 7 else 7.8
    _set_repeat_table_header(table.rows[0])
    for index, (header, width) in enumerate(zip(headers, widths, strict=True)):
        cell = table.rows[0].cells[index]
        _set_table_cell_width(cell, width)
        _set_cell_shading(cell, LIGHT_GRAY)
        _set_cell_margins(cell, top=80, start=80, bottom=80, end=80)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = header
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        _set_run_font(paragraph.runs[0], font_size, bold=True, color=DARK_BLUE)
    for row_index, values in enumerate(body):
        row = table.add_row()
        _prevent_row_split(row)
        for column_index, (value, width) in enumerate(zip(values, widths, strict=True)):
            cell = row.cells[column_index]
            _set_table_cell_width(cell, width)
            _set_cell_margins(cell, top=70, start=75, bottom=70, end=75)
            if row_index % 2:
                _set_cell_shading(cell, "FAFBFC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = value
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _set_run_font(paragraph.runs[0], font_size)

    note = document.add_paragraph(style="Manuscript Note")
    role = ", ".join(cast("Any", item).evidence_roles)
    run = note.add_run(f"Evidence role: {role}. Source: checksum-verified publication display.")
    _set_run_font(run, 8, color=GRAY)


def _set_picture_alt_text(run: Run, *, title: str, description: str) -> None:
    doc_pr = run._r.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("title", title)
        doc_pr[0].set("descr", description)


def _add_figure(document: DocumentObject, display_path: Path, item: object) -> None:
    item_id = cast("Any", item).item_id
    image_path = display_path / next(
        file for file in cast("Any", item).files if file.endswith(".png")
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    _set_picture_alt_text(
        run,
        title=cast("Any", item).title,
        description=cast("Any", item).caption,
    )
    caption = document.add_paragraph(style="Manuscript Caption")
    caption.paragraph_format.keep_with_next = False
    first = caption.add_run(f"{_figure_label(item_id)}. {cast('Any', item).title}. ")
    _set_run_font(first, 9, bold=True)
    second = caption.add_run(cast("Any", item).caption)
    _set_run_font(second, 9)


def _add_references(
    document: DocumentObject,
    ordered_references: tuple[ManuscriptReference, ...],
) -> None:
    for index, reference in enumerate(ordered_references, start=1):
        paragraph = document.add_paragraph(style="Reference Entry")
        run = paragraph.add_run(f"{index}. {reference.formatted} {reference.persistent_url}")
        _set_run_font(run, 9)


def _iter_blocks(source: str) -> list[str]:
    blocks: list[str] = []
    paragraph_lines: list[str] = []
    for line in source.splitlines():
        stripped = line.strip()
        is_structural = (
            not stripped
            or stripped.startswith("#")
            or stripped.startswith("{{")
            or stripped.startswith("- ")
        )
        if is_structural and paragraph_lines:
            blocks.append(" ".join(paragraph_lines))
            paragraph_lines = []
        if not stripped:
            continue
        if stripped.startswith("#") or stripped.startswith("{{") or stripped.startswith("- "):
            blocks.append(stripped)
        else:
            paragraph_lines.append(stripped)
    if paragraph_lines:
        blocks.append(" ".join(paragraph_lines))
    return blocks


def _render_docx(
    spec: ManuscriptSpec,
    manuscript_source: str,
    ordered_references: tuple[ManuscriptReference, ...],
) -> bytes:
    inventory, _ = read_publication_display(spec.display.path, require_publication_ready=True)
    item_index = {item.item_id: item for item in inventory.items}
    cited_source = replace_citations(manuscript_source, ordered_references)
    document = Document()
    _configure_document(document, spec)
    _add_cover(document, spec, manuscript_source)
    for block in _iter_blocks(cited_source):
        if block == "{{pagebreak}}":
            _add_page_break(document)
            continue
        if block == "{{references}}":
            _add_references(document, ordered_references)
            continue
        marker = _MARKER_PATTERN.fullmatch(block)
        if marker:
            kind, item_id = marker.groups()
            item = item_index[item_id]
            if kind == "table":
                _add_table(document, spec.display.path, item)
            else:
                _add_figure(document, spec.display.path, item)
            continue
        if block.startswith("### "):
            document.add_heading(block[4:], level=3)
        elif block.startswith("## "):
            document.add_heading(block[3:], level=2)
        elif block.startswith("# "):
            document.add_heading(block[2:], level=1)
        elif block.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, block[2:])
        else:
            _add_body_paragraph(document, block)
    document.sections[-1].start_type = WD_SECTION_START.NEW_PAGE
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _assemble_markdown(
    spec: ManuscriptSpec,
    manuscript_source: str,
    ordered_references: tuple[ManuscriptReference, ...],
) -> str:
    inventory, _ = read_publication_display(spec.display.path, require_publication_ready=True)
    item_index = {item.item_id: item for item in inventory.items}
    rendered = replace_citations(manuscript_source, ordered_references)
    rendered = rendered.replace("{{pagebreak}}", "")
    reference_lines = [
        f"{index}. {reference.formatted} {reference.persistent_url}"
        for index, reference in enumerate(ordered_references, start=1)
    ]
    rendered = rendered.replace("{{references}}", "\n".join(reference_lines))
    for item_id in spec.included_tables:
        item = item_index[item_id]
        markdown_path = spec.display.path / next(
            file for file in item.files if file.endswith(".md")
        )
        table_markdown = markdown_path.read_text(encoding="utf-8").strip()
        rendered = rendered.replace(f"{{{{table:{item_id}}}}}", table_markdown)
    for item_id in spec.included_figures:
        item = item_index[item_id]
        image_path = "../paper-display-v1/" + next(
            file for file in item.files if file.endswith(".png")
        )
        figure_markdown = (
            f"![{item.title}]({image_path})\n\n"
            f"**{_figure_label(item_id)}. {item.title}.** {item.caption}"
        )
        rendered = rendered.replace(f"{{{{figure:{item_id}}}}}", figure_markdown)
    return rendered.strip() + "\n"


def _sync_or_verify_latex_sources(
    spec: ManuscriptSpec,
    package: LatexManuscript,
    *,
    sync: bool,
) -> None:
    expected_text = {
        spec.latex_source: package.source,
        spec.latex_bibliography: package.bibliography,
    }
    expected_binary = {
        spec.latex_source.parent / relative_name: content
        for relative_name, content in package.figures.items()
    }
    if sync:
        for path, text_content in expected_text.items():
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(text_content, encoding="utf-8")
        for path, binary_content in expected_binary.items():
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(binary_content)
        return
    stale = [
        str(path)
        for path, content in expected_text.items()
        if not path.is_file() or path.read_text(encoding="utf-8") != content
    ]
    stale.extend(
        str(path)
        for path, content in expected_binary.items()
        if not path.is_file() or path.read_bytes() != content
    )
    if stale:
        raise RuntimeError(
            "tracked LaTeX sources are absent or stale: "
            + ", ".join(stale)
            + "; rerun with --sync-latex-source --allow-dirty, inspect, and commit them"
        )


def _compile_latex(spec: ManuscriptSpec, package: LatexManuscript) -> bytes:
    tectonic = shutil.which("tectonic")
    if tectonic is None:
        raise RuntimeError(
            "Tectonic is required to compile the verified LaTeX manuscript; "
            "install it with `brew install tectonic`"
        )
    latex_name = spec.latex_source.name
    with tempfile.TemporaryDirectory(prefix="neuroselect-latex-") as temporary:
        workspace = Path(temporary)
        (workspace / latex_name).write_text(package.source, encoding="utf-8")
        (workspace / spec.latex_bibliography.name).write_text(
            package.bibliography,
            encoding="utf-8",
        )
        for relative_name, content in package.figures.items():
            target = workspace / relative_name
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(content)
        completed = subprocess.run(
            [
                tectonic,
                "--keep-logs",
                "--outdir",
                str(workspace),
                latex_name,
            ],
            cwd=workspace,
            capture_output=True,
            text=True,
        )
        if completed.returncode:
            diagnostic = "\n".join(
                line
                for line in (completed.stdout + "\n" + completed.stderr).splitlines()[-60:]
                if line.strip()
            )
            raise RuntimeError(f"LaTeX compilation failed:\n{diagnostic}")
        pdf_path = workspace / Path(latex_name).with_suffix(".pdf")
        if not pdf_path.is_file():
            raise RuntimeError("Tectonic completed without producing the manuscript PDF")
        pdf_bytes = pdf_path.read_bytes()
    if not pdf_bytes.startswith(b"%PDF"):
        raise RuntimeError("Tectonic output is not a PDF document")
    return pdf_bytes


def _latex_readme(spec: ManuscriptSpec) -> bytes:
    return (
        "# NeuroSelect LaTeX manuscript\n\n"
        "Compile the self-contained bundle with:\n\n"
        "```bash\n"
        f"tectonic {spec.latex_source.name}\n"
        "```\n\n"
        "The committed `paper/latex` source is generated from the verified Markdown, "
        "reference registry, publication tables, and publication figures. Regenerate it with "
        '`make manuscript MANUSCRIPT_ARGS="--sync-latex-source --allow-dirty --overwrite"` '
        "after any manuscript-source change.\n"
    ).encode()


def main() -> None:
    args = parse_args()
    spec = load_manuscript_spec(args.config)
    (
        manuscript_source,
        ordered_references,
        claim_audit,
        display_inventory,
        display_manifest,
    ) = verify_manuscript_inputs(spec)
    latex_package = render_latex_manuscript(
        spec,
        manuscript_source,
        ordered_references,
        display_inventory,
    )
    _sync_or_verify_latex_sources(
        spec,
        latex_package,
        sync=args.sync_latex_source,
    )
    revision, source_tree_sha256 = git_state()
    if source_tree_sha256 is not None and not args.allow_dirty:
        raise RuntimeError(
            "refusing publication manuscript assembly from a dirty worktree; commit changes or "
            "use --allow-dirty for development-only visual QA"
        )
    document_bytes = _render_docx(spec, manuscript_source, ordered_references)
    rendered_markdown = _assemble_markdown(spec, manuscript_source, ordered_references)
    latex_pdf = _compile_latex(spec, latex_package)
    latex_pdf_name = Path(spec.latex_source.name).with_suffix(".pdf").name
    latex_payloads = {
        spec.latex_source.name: latex_package.source.encode(),
        spec.latex_bibliography.name: latex_package.bibliography.encode(),
        latex_pdf_name: latex_pdf,
        "LATEX-README.md": _latex_readme(spec),
        **latex_package.figures,
    }
    inventory, manifest = write_manuscript_bundle(
        spec,
        rendered_markdown=rendered_markdown,
        document_bytes=document_bytes,
        additional_files=latex_payloads,
        claim_audit=claim_audit,
        ordered_references=ordered_references,
        display_inventory=display_inventory,
        display_manifest=display_manifest,
        output_dir=args.output,
        git_sha=revision,
        source_tree_sha256=source_tree_sha256,
        overwrite=args.overwrite,
    )
    print(f"Run: {manifest.run_id}")
    print(f"Manuscript: {args.output / spec.output_filename}")
    print(f"LaTeX: {args.output / spec.latex_source.name}")
    print(f"Compiled PDF: {args.output / latex_pdf_name}")
    print(f"Citations: {inventory.citation_count}")
    print(f"Verified quantitative claims: {inventory.quantitative_claim_count}")
    print(f"Tables / figures: {len(inventory.included_tables)} / {len(inventory.included_figures)}")
    print(f"Assembly ready: {inventory.assembly_ready}")
    print(f"Submission ready: {inventory.submission_ready}")
    print(f"Pending external gates: {', '.join(inventory.pending_submission_gates)}")
    print(f"Manifest SHA-256: {manifest.digest()}")


if __name__ == "__main__":
    main()
